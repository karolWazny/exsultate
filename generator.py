import os
import time
from exsultate import *
from docx import Document

log_filename = os.path.dirname(os.path.realpath(__file__)) + '/log.txt'

class Generator:
    def __init__(self):
        self.path = os.path.dirname(os.path.realpath(__file__)) + '/'
        self.configuration = None
        self.filename = None
        self.document = None

    def load_configuration(self, filename="config.json"):
        if not os.path.isfile(self.path + filename):
            self.generate_config_file(self.path + filename)
        self.configuration = read_json(self.path + filename)
        
    @staticmethod
    def get_default_config():
        return {"style_mappings": { "number": "number", "verse": "verse", "chorus": "chorus", "title": "title-itself"},
                "template": "template.docx"}

    def generate_config_file(self, filename):
        write_json(Generator.get_default_config(), filename)

    def generate(self, songbook, filename = ''):
        # self.filename = filename + str(time.time_ns()) + ".docx"
        self.filename = filename + ".docx"
        self.document = Document(self.path + self.configuration['template'])
        for song in songbook.songs:
            self.add_song_to_document(song)

        self.document.save(self.path + self.filename)
        return self.filename

    def add_song_to_document(self, song):
        self.document.add_paragraph('', style=self.configuration['style_mappings']['number'])
        self.document.add_paragraph(song.title, style=self.configuration['style_mappings']['title'])
        songparts = song.content
        for songpart in songparts:
            self.add_songpart_to_document(songpart)



    def add_songpart_to_document(self, songpart):
        self.document.add_paragraph(songpart.lyrics(), style=self.configuration['style_mappings'][songpart.type])